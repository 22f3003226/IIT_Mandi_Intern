from pathlib import Path
from typing import Optional

import docx

from app.parsers.base import looks_like_equation
from app.schemas.parsed_document import (
    DocumentMetadata,
    EquationRef,
    FigureRef,
    ParsedDocument,
    Section,
    TableBlock,
)

HEADING_STYLES = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3, "Title": 1}


def _flush(sections: list[Section], buffer: list[str], heading: Optional[str], level: int) -> None:
    if buffer:
        sections.append(Section(heading=heading, level=level, text="\n".join(buffer).strip()))


def parse_docx(file_path: str) -> ParsedDocument:
    document = docx.Document(file_path)
    sections: list[Section] = []
    equations: list[EquationRef] = []
    current_heading: Optional[str] = None
    current_level = 0
    buffer: list[str] = []

    for para in document.paragraphs:
        text = para.text
        if not text.strip():
            continue
        if looks_like_equation(text):
            equations.append(EquationRef(text=text.strip()))
        style_name = para.style.name if para.style else ""
        if style_name in HEADING_STYLES:
            _flush(sections, buffer, current_heading, current_level)
            buffer = []
            current_heading = text.strip()
            current_level = HEADING_STYLES[style_name]
        else:
            buffer.append(text)
    _flush(sections, buffer, current_heading, current_level)

    tables = [
        TableBlock(rows=[[cell.text for cell in row.cells] for row in table.rows])
        for table in document.tables
    ]
    figures = [FigureRef() for _ in document.inline_shapes]

    return ParsedDocument(
        metadata=DocumentMetadata(source_filename=Path(file_path).name, format="docx", page_count=1),
        sections=sections,
        tables=tables,
        figures=figures,
        equations=equations,
    )
